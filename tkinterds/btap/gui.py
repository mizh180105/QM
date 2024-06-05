from pathlib import Path
from tkinter import Canvas, Text, Button, PhotoImage, font, filedialog, messagebox, Scrollbar
import tkinter as tk
from docx import Document
import random
import string
OUTPUT_PATH = Path(__file__).parent
ASSETS_PATH = OUTPUT_PATH / Path(r"D:\tkinterds\build\assets\frame0")
window = tk.Tk()
window.title("Question mixing")
window.geometry("1300x800")
window.configure(bg = "#D1D1D1")
font_size = tk.StringVar()

textbox = Text(
        bd=1,  # Set the border width to 1
        bg="#FFFFFF",
        fg="#000716",
        highlightthickness=1,
        relief="solid",
        padx=5,
        pady=5
)
textbox.place(
        x=90.0,
        y=155.0,
        width=1110.0,
        height=1000.0  # Modify the height to make it infinite
)
class questionmixer():
        def relative_to_assets(path: str) -> Path:
                return ASSETS_PATH / Path(path)
        def make_text_bold():
                current_font = font.Font(font=textbox['font'])
                if current_font.actual()['weight'] == 'bold':
                        current_font.configure(weight='normal')
                else:
                        current_font.configure(weight='bold')
                textbox.configure(font=current_font)
        def make_text_italic():
                current_font = font.Font(font=textbox['font'])
                if current_font.actual()['slant'] == 'italic':
                        current_font.configure(slant='roman')
                else:
                        current_font.configure(slant='italic')
                textbox.configure(font=current_font)
        def make_text_underline():
                current_font = font.Font(font=textbox['font'])
                if current_font.actual()['underline']:
                        current_font.configure(underline=False)
                else:
                        current_font.configure(underline=True)
                textbox.configure(font=current_font)
        def make_text_subscript():
                current_font = font.Font(font=textbox['font'])
                current_font.configure(size=8)
                textbox.configure(font=current_font, spacing3=5)
        def make_text_superscript():
                current_font = font.Font(font=textbox['font'])
                current_font.configure(size=8)
                textbox.configure(font=current_font, spacing1=-5)
        def align_text_left():
                textbox.tag_configure("left", justify='left')
                textbox.tag_add("left", 1.0, "end")
        def align_text_center():
                textbox.tag_configure("center", justify='center')
                textbox.tag_add("center", 1.0, "end")
        def align_text_right():
                textbox.tag_configure("right", justify='right')
                textbox.tag_add("right", 1.0, "end")
        def insert_pi():
                textbox.insert("end", "π")
        def insert_infinity():
                textbox.insert("end", "∞")
        def insert_sigma():
                textbox.insert("end", "Σ")
        def insert_integral():
                textbox.insert("end", "∫")
        def insert_sqrt():
                textbox.insert("end", "√")
        def insert_approx():
                textbox.insert("end", "≈")
        def insert_alpha():
                textbox.insert("end", "α")
        def insert_beta():
                textbox.insert("end", "β")
        def insert_delta():
                textbox.insert("end", "δ")
        def insert_leq():
                textbox.insert("end", "≤")
        def insert_geq():
                textbox.insert("end", "≥")
        def insert_neq():
                textbox.insert("end", "≠")
        def open_word_file():
                file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
                if file_path:
                        doc = Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                        full_text.append(para.text)
                textbox.insert('1.0', '\n'.join(full_text))
        def mix_questions_and_answers():
                content = textbox.get('1.0', 'end').strip().split('\n')
                questions = [line for line in content if line.startswith('Q:')]
                answers = [line for line in content if line.startswith('A:')]
                textbox.tag_config('question', foreground='blue')
                textbox.tag_config('answer', foreground='green')
                if len(questions) != len(answers):
                        messagebox.showerror("Error")
                        return
                qa_pairs = list(zip(answers, questions))  # Swap the positions of questions and answers
                random.shuffle(qa_pairs)
                shuffled_answers, shuffled_questions = zip(*qa_pairs)  # Swap the positions of questions and answers
                textbox.delete('1.0', 'end')
                test_code = ''.join(random.choices(string.digits, k=4))
                textbox.insert('end', f'Test Code: {test_code}\n')
                for i in range(len(shuffled_questions)):
                        textbox.insert('end', f'Q{i+1}: {shuffled_questions[i]}\n', 'question')
                        answer_options = shuffled_answers[i].split(',')
                        for j in range(len(answer_options)):
                                textbox.insert('end', f'{chr(65+j)}: {answer_options[j]}\n', 'answer')

        textbox.tag_config('question', foreground='blue')
        textbox.tag_config('answer', foreground='green')
        def export_to_word():
                filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
                if filename:
                        content = textbox.get('1.0', 'end').strip().split('\n')
                        doc = Document()
                        for line in content:
                                doc.add_paragraph(line)
                        doc.save(filename)
                        messagebox.showinfo("Export Successful", "File exported successfully.")
        canvas = Canvas(
                window,
                bg="#D1D1D1",
                height=120,
                width=1300,
                bd=0,
                highlightthickness=0,
                relief="ridge"
        )
        
       

        canvas.place(x=0, y=0)

        canvas.create_rectangle(
                0.0,
                0.0,
                1300.0,
                110.0,
                fill="#D3D3D3",
                outline="#5F5F5F",
                width=1.0
        )

        canvas.create_rectangle(
                420.0,
                60.0,
                425.0,
                90.0,
                fill="#979797",
                outline="")
        canvas.create_rectangle(
                600.0,
                60.0,
                605.0,
                90.0,
                fill="#979797",
                outline="")

        canvas.create_rectangle(
                775.0,
                60.0,
                775.0,
                90.0,
                fill="#979797",
                outline="")

        canvas.create_text(
                190.0,
                95.0,
                anchor="nw",
                text="Phông chữ",
                fill="#000000",
                font=("SegoeUI Semilight", 12 * -1)
                )

        canvas.create_text(
                490.0,
                95.0,
                anchor="nw",
                text="Căn chỉnh",
                fill="#000000",
                font=("SegoeUI Semilight", 12 * -1)
                )

        canvas.create_text(
                651.0,
                94.0,
                anchor="nw",
                text="Ký tự đặc biệt",
                fill="#000000",
                font=("SegoeUI Semilight", 12 * -1)
                )
        textbox.insert('1.0', '\u200e')
        textbox.delete('1.0', 'end')
        textbox.tag_configure("margin", lmargin1=50, rmargin=50)  # Set left and right margins to 50 pixels
        textbox.tag_add("margin", "1.0", "end")  # Apply the tag to the entire text
        textbox.config(highlightbackground="#000000", highlightcolor="#000000")

        def make_adjust_font(font_name):
                def adjust_font():
                        textbox['font'] = font.Font(family=font_name)
                return adjust_font
        def make_adjust_font_size(size):
                def adjust_font_size():
                        textbox['font'] = font.Font(size=int(size))
                return adjust_font_size

        font_name = tk.StringVar()  
        font_size = tk.StringVar()  # Assuming font_size is a StringVar

        menu_button = tk.Menubutton(window, text="                 >", relief="raised")
        menu_button.place(x=20.0, y=64.0, width=80.0, height=20.0)

        menu = tk.Menu(menu_button, tearoff=False)
        menu_button["menu"] = menu

        fonts = ['Helvetica', 'Times', 'Arial', 'SegoeUI','Inter']  # Add more fonts as needed

        for f in fonts:
                menu.add_radiobutton(label=f, variable=font_name, command=make_adjust_font(f))

        entry = tk.Entry(window, textvariable=font_name)
        entry.place(x=20.0, y=64.0, width=60.0, height=20.0)  # Changed y position to avoid overlap

        menu_button = tk.Menubutton(window, text="     >", relief="raised")
        menu_button.place(x=190.0, y=65.0, width=40.0, height=20.0)

        menu = tk.Menu(menu_button, tearoff=False)
        menu_button["menu"] = menu

        for size in range(10, 60, 2):  
                menu.add_radiobutton(label=str(size), variable=font_size, command=make_adjust_font_size(size))

        entry = tk.Entry(window, textvariable=font_size)
        entry.place(x=190.0, y=64.0, width=20.0, height=20.0)  # Changed y position to avoid overlap
        canvas.create_rectangle(
                0.0,
                0.0,
                1300.0,
                51.0,
                fill="#D9D9D9",
                outline="#5F5F5F",
                width=0.5)
        button_image_1 = PhotoImage(
                file=relative_to_assets("button_1.png"))
        button_1 = Button(
                window,
                image=button_image_1,
                borderwidth=0,
                highlightthickness=0,
                command= make_text_bold,
                relief="flat"
                        )
        button_1.place(
                x=247.0,
                y=65.0,
                width=20.0,
                height=20.0
                        )
        button_image_2 = PhotoImage(
                file=relative_to_assets("button_2.png"))
        button_2 = Button(
                image=button_image_2,
                borderwidth=0,
                highlightthickness=0,
                command=make_text_italic,
                relief="flat"
                )
        button_2.place(
                x=282.0,
                y=65.0,
                width=20.0,
                height=20.0
                )

        button_image_3 = PhotoImage(
                file=relative_to_assets("button_3.png"))
        button_3 = Button(
                image=button_image_3,
                borderwidth=0,
                highlightthickness=0,
                command=make_text_underline,
                relief="flat"
                        )
        button_3.place(
                x=317.0,
                y=65.0,
                width=20.0,
                height=20.0
                        )
        button_image_4 = PhotoImage(
                        file=relative_to_assets("button_4.png"))
        button_4 = Button(
                image=button_image_4,
                borderwidth=0,
                highlightthickness=0,
                command=open_word_file,
                relief="flat"
                )
        button_4.place(
                        x=20.0,
                        y=17.0,
                        width=88.0,
                        height=19.0
                        )

        button_image_5 = PhotoImage(
                        file=relative_to_assets("button_5.png"))
        button_5 = Button(
                image=button_image_5,
                borderwidth=0,
                highlightthickness=0,
                command=mix_questions_and_answers,
                relief="flat"
                        )
        button_5.place(
                x=131.0,
                y=17.0,
                width=88.0,
                height=19.0
                )

        button_image_6 = PhotoImage(
                file=relative_to_assets("button_6.png"))
        button_6 = Button(
                image=button_image_6,
                borderwidth=0,
                highlightthickness=0,
                command=export_to_word,
                relief="flat"
                )
        button_6.place(
                x=242.0,
                y=17.0,
                width=88.0,
                height=19.0
                )

        button_image_7 = PhotoImage(
                file=relative_to_assets("button_7.png"))
        button_7 = Button(
                image=button_image_7,
                borderwidth=0,
                highlightthickness=0,
                command=make_text_subscript,
                relief="flat"
                        )
        button_7.place(
                x=349.0,
                y=67.0,
                width=21.0,
                height=20.0
                        )

        button_image_8 = PhotoImage(
                file=relative_to_assets("button_8.png"))
        button_8 = Button(
                image=button_image_8,
                borderwidth=0,
                highlightthickness=0,
                command=make_text_superscript,
                relief="flat"
                        )
        button_8.place(
                x=381.0,
                y=65.0,
                width=21.0,
                height=20.0
                        )

        button_image_9 = PhotoImage(
                file=relative_to_assets("button_9.png"))
        button_9 = Button(
                image=button_image_9,
                borderwidth=0,
                highlightthickness=0,
                command=align_text_left,
                relief="flat"
                        )
        button_9.place(
                x=438.0,
                y=65.0,
                width=21.0,
                height=20.0
                        )

        button_image_10 = PhotoImage(
                file=relative_to_assets("button_10.png"))
        button_10 = Button(
                image=button_image_10,
                borderwidth=0,
                highlightthickness=0,
                command=align_text_center,
                relief="flat"
                        )
        button_10.place(
                x=500.0,
                y=65.0,
                width=21.0,
                height=20.0
                        )

        button_image_11 = PhotoImage(
                file=relative_to_assets("button_11.png"))
        button_11 = Button(
                image=button_image_11,
                borderwidth=0,
                highlightthickness=0,
                command= align_text_right,
                relief="flat"
                        )
        button_11.place(
                x=560.0,
                y=65.0,
                width=21.0,
                height=20.0
                        )

        button_image_13 = PhotoImage(
                file=relative_to_assets("button_13.png"))
        button_13 = Button(
                image=button_image_13,
                borderwidth=0,
                highlightthickness=0,
                command=insert_pi,
                relief="flat"
                        )
        button_13.place(
                x=614.0,
                y=52.0,
                width=19.0,
                height=20.0
                        )

        button_image_14 = PhotoImage(
                file=relative_to_assets("button_14.png"))
        button_14 = Button(
                image=button_image_14,
                borderwidth=0,
                highlightthickness=0,
                command=insert_infinity,
                relief="flat"
                        )
        button_14.place(
                x=640.0,
                y=52.0,
                width=19.0,
                height=17.0
                )

        button_image_15 = PhotoImage(
                file=relative_to_assets("button_15.png"))
        button_15 = Button(
                image=button_image_15,
                borderwidth=0,
                highlightthickness=0,
                command=insert_sigma,
                relief="flat"
                        )
        button_15.place(
                x=666.0,
                y=52.0,
                width=19.0,
                height=19.0
                        )

        button_image_16 = PhotoImage(
                file=relative_to_assets("button_16.png"))
        button_16 = Button(
                image=button_image_16,
                borderwidth=0,
                highlightthickness=0,
                command=insert_sqrt,
                relief="flat"
                        )
        button_16.place(
                x=713.0,
                y=53.0,
                width=19.0,
                height=18.0
                        )

        button_image_17 = PhotoImage(
                file=relative_to_assets("button_17.png"))
        button_17 = Button(
                image=button_image_17,
                borderwidth=0,
                highlightthickness=0,
                command=insert_approx,
                relief="flat"
                        )
        button_17.place(
                x=738.0,
                y=52.0,
                width=19.0,
                height=19.0
                        )

        button_image_18 = PhotoImage(
                file=relative_to_assets("button_18.png"))
        button_18 = Button(
                image=button_image_18,
                borderwidth=0,
                highlightthickness=0,
                command=insert_integral,
                relief="flat"
                        )
        button_18.place(
                x=690.0,
                y=53.0,
                width=19.0,
                height=18.0
                        )

        button_image_19 = PhotoImage(
                file=relative_to_assets("button_19.png"))
        button_19 = Button(
                image=button_image_19,
                borderwidth=0,
                highlightthickness=0,
                command=insert_alpha,
                relief="flat"
                        )
        button_19.place(
                x=614.0,
                y=71.0,
                width=19.0,
                height=19.0
                        )

        button_image_20 = PhotoImage(
                file=relative_to_assets("button_20.png"))
        button_20 = Button(
                image=button_image_20,
                borderwidth=0,
                highlightthickness=0,
                command=insert_beta,
                relief="flat"
                        )
        button_20.place(
                x=640.0,
                y=71.0,
                width=19.0,
                height=19.0
                        )

        button_image_21 = PhotoImage(
                file=relative_to_assets("button_21.png"))
        button_21 = Button(
                image=button_image_21,
                borderwidth=0,
                highlightthickness=0,
                command=insert_delta,
                relief="flat"
                        )
        button_21.place(
                x=666.0,
                y=71.0,
                width=19.0,
                height=18.0
                        )

        button_image_22 = PhotoImage(
                file=relative_to_assets("button_22.png"))
        button_22 = Button(
                image=button_image_22,
                borderwidth=0,
                highlightthickness=0,
                command=insert_leq,
                relief="flat"
                        )
        button_22.place(
                x=689.0,
                y=71.0,
                width=19.0,
                height=18.0
                        )

        button_image_23 = PhotoImage(
                file=relative_to_assets("button_23.png"))
        button_23 = Button(
                image=button_image_23,
                borderwidth=0,
                highlightthickness=0,
                command= insert_geq,
                relief="flat"
                        )
        button_23.place(
                x=714.0,
                y=71.0,
                width=19.0,
                height=18.0
                        )

        button_image_24 = PhotoImage(
                file=relative_to_assets("button_24.png"))
        button_24 = Button(
                image=button_image_24,
                borderwidth=0,
                highlightthickness=0,
                command=insert_neq,
                relief="flat"
                        )
        button_24.place(
                x=738.0,
                y=71.0,
                        width=19.0,
                        height=19.0
                        )

        window.resizable(False, False)
        window.mainloop()  
